from docx import Document
from docx.shared import Pt

doc = Document()  # Create the doc

rows = [["Page Name", "Content Creator", "Approver", "Word Count", "Live URL", "New URL",  "Page Title", "Meta Description", "Live URL", 
        "New URL"], ["SEO Data:", "Page Title (40-60 Characters)", "Meta Description (140-160 Characters)", "Focus Keyword", "Key Variations",
        "Sprinkle Keywords (Use if you can, not required)", "Misc/Notes"], ["Onpage:", "Content"]
        ]
sections = ["Page Info:", "SEO Data:", "Onpage:"]
keys = ["Page Name", "Page Title", "Meta Description", "Live URL"]  # The Rows we need to fill out (except for content)
whitelist = set('abcdefghijklmnopqrstuvwxyz ABCDEFGHIJKLMNOPQRSTUVWXYZ')  # Used to filter garbage when reading rows

doc.add_picture('img/logo.png')  # Add the Astoundz img to the top of the doc

table = doc.add_table(rows=1, cols=2)

row = table.rows[0].cells

# Add content to table
row[0].text = "Page Info:"  # Initial setup
for i in rows:
    for item in i:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = "Test"

# Format table font
for row in table.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.name = "Arial"
                font.size= Pt(11)
                paragraph.paragraph_format.line_spacing = 1.75 if cell.text not in rows else 0



table.style = 'Table Grid'

# for table in doc.tables:
#     for row in table.rows:
#         cells = row.cells
#         for cell in range(len(cells)):

#             cellContent = ''.join(filter(whitelist.__contains__, cells[cell].text.split("\n")[0]))  # Get current cell content

#             if cellContent in keys:
#                 cells[cell].text = "Test"

doc.save("test.docx")  # Save the doc