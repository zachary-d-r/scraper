from copy import deepcopy
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup

class document:

    keys = ["Page Name", "Page Title", "Meta Description", "Live URL", "Authority Score", "Backlinks"]  # The Rows we need to fill out (except for content)
    whitelist = set('abcdefghijklmnopqrstuvwxyz ABCDEFGHIJKLMNOPQRSTUVWXYZ')  # Used to filter garbage when reading rows

    def __init__(self, numOfTables):
        self.doc = Document()  # Create the doc

        self.doc.add_picture('img/logo.png')  # Add the Astoundz img to the top of the doc

        for i in range(numOfTables): self.createTable()  # Create all the tables we need



    def getTable(self):
        """
        Return the table we are going to fill out by copying it from another document

        Returns:
        ========
            A copy of the table from table.docx
        """     
        tableDoc = Document("imports/table.docx")  # Import the doc with the table
        return deepcopy(tableDoc.tables[0]._tbl)  # Return a copy of that table

    def createTable(self):
        """
        Create a table that gets put in the doc
        """
        paragraph = self.doc.add_paragraph()  # Add a paragraph to seperate the tables
        paragraph._p.addnext(self.getTable())  # Put the table in the doc
        
        # Add a newline after the table is copied
        r = paragraph.add_run()
        r.add_break()

    def insertContent(self, content):
        """
        Insert the first few rows of content
        """

        # Go through each table
        for table in range(len(self.doc.tables)):
            contentNum = 0  # This variable is used to determine which URL we should be putting info in for

            # For every row in each cell...
            for row in self.doc.tables[table].rows:
                for cell in range(len(row.cells)):
                    cellContent = ''.join(filter(self.whitelist.__contains__, row.cells[cell].text.split("\n")[0]))  # Get current cell content
                    
                    # Check if the cell is a row we want to fill out
                    if cellContent in self.keys: 
                        row.cells[cell+1].text = content[table][contentNum]  # If it is, replace the text

                        # Style the text
                        for p in row.cells[cell+1 if cell+1 < len(row.cells) else cell].paragraphs:
                            for run in p.runs:
                                font = run.font  # Create a font
                                font.name = "Arial"  # Set the font to arial
                                font.size= Pt(11) if cellContent != "Page Name" else Pt(18)  # Set the font size
                                font.bold = False if cellContent != "Page Name" else True  # Set if the font should be bold or not

                        contentNum += 1

    def insertFormatedContent(self, content):
        """
        Insert the formatted content
        """

        #print(len(content))
        
        i = 0  # Iterator used for what table we are on
        for table in self.doc.tables:

            cell = table.rows[len(table.rows)-1].cells[1]  # Get the cell we want to put our content in

            # For every item in the content list...
            for item in content[i if i < len(content) else len(content)-1]:
          
                #print(cell.text)

                #print(tag.name + ' -> ' + tag.text.strip())

                tagName = item.name  # Get the tag name of the item
                t = item.text  # Get the innerhtml of the item
                text = ""

                l = t.split("  ")
                for j in l:
                    text += j


                # Insert and format text based on tag name
                if tagName == "h2":
                    p = cell.add_paragraph()
                    run = p.add_run(text)
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(20)

                elif tagName == "h3":
                    p = cell.add_paragraph()
                    run = p.add_run(text)
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(16)

                elif tagName == "p":
                    p = cell.add_paragraph()
                    run = p.add_run(text)
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(11)

                
                #for p in row.cells[cell+1 if cell+1 < len(row.cells) else cell].paragraphs:

            i += 1  # Increase iterator
                    

    def save(self, filename):
        self.doc.save(filename)  # Save the doc