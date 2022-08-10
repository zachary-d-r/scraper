from copy import deepcopy
from docx import Document
from docx.shared import Pt
import imports.Scrape as Scrape

def getTable():
    """
    Return the table we are going to fill out
    """
    tableDoc = Document("imports/table.docx")
    return deepcopy(tableDoc.tables[0]._tbl)


doc = Document()  # Create the doc
scraper = Scrape.scrape(r"https://github.com/")  # Create scraper object
s = Scrape.scrape(r"https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html")

keys = ["Page Name", "Page Title", "Meta Description", "Live URL"]  # The Rows we need to fill out (except for content)
whitelist = set('abcdefghijklmnopqrstuvwxyz ABCDEFGHIJKLMNOPQRSTUVWXYZ')  # Used to filter garbage when reading rows

doc.add_picture('img/logo.png')  # Add the Astoundz img to the top of the doc

# Create the table
paragraph = doc.add_paragraph()
paragraph._p.addnext(getTable())
r = paragraph.add_run()
r.add_break()
r.add_break()
paragraph = doc.add_paragraph()
paragraph._p.addnext(getTable())

def insertContent(content):
    for table in range(len(doc.tables)):
        contentNum = 0
        for row in doc.tables[table].rows:
            for cell in range(len(row.cells)):
                cellContent = ''.join(filter(whitelist.__contains__, row.cells[cell].text.split("\n")[0]))  # Get current cell content
                if cellContent in keys: 
                    row.cells[cell+1].text = content[table][contentNum]

            
                    for p in row.cells[cell+1 if cell+1 < len(row.cells) else cell].paragraphs:
                        for run in p.runs:
                            font = run.font
                            font.name = "Arial"
                            font.size= Pt(11) if cellContent != "Page Name" else Pt(18)
                            font.bold = False if cellContent != "Page Name" else True

                    contentNum += 1

insertContent([[scraper.h1, scraper.URL, scraper.title, scraper.metaDescription], [s.h1, s.URL, s.title, s.metaDescription]])

doc.save("test.docx")  # Save the doc